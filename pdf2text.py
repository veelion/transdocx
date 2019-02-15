#!/usr/bin/env python
# coding:utf-8
# Author: veelion@ebuinfo.com

import sys
import os
import subprocess
from docx import Document


def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base_path, relative_path)


if os.name == 'nt':
    exe_bin = 'bin/pdftotext.exe'
else:
    exe_bin = 'bin/pdftotext'

exe_bin = resource_path(exe_bin)
print('exe_bin:', exe_bin)


def pdf_to_text(fn):
    fn_text = 'tmp-pdf.txt'
    args = [
        exe_bin,
        '-enc', 'UTF-8',
        fn,
        fn_text
    ]
    print(args)

    p = subprocess.Popen(args,
                         shell=True,
                         stdin=subprocess.PIPE,
                         stdout=subprocess.PIPE,
                         stderr=subprocess.PIPE)
    out, err = p.communicate()
    print('out:', out.decode())
    print('err:', err.decode())
    try:
        text = open(fn_text, 'rb').read()
        text = text.decode('utf8', 'ignore')
    except Exception as e:
        print(e)
        text = ''
    return text


def pdf_to_docx(fn):
    text = pdf_to_text(fn)
    print('text:', len(text))
    text = text.replace('\x0c', '')
    lines = text.split('\n')
    doc = Document(resource_path('./bin/default.docx'))
    for line in lines:
        line = line.strip()
        if not line: continue
        doc.add_paragraph(line)
    fn_docx = fn + '.docx'
    print('paragraphs:', len(doc.paragraphs))
    doc.save(fn_docx)
    return fn_docx


if __name__ == '__main__':
    import sys
    fn = sys.argv[1]
    text = pdf_to_docx(fn)
    print(text)
